import streamlit as st
from docx import Document
import re
import io
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 変換ロジック（前述の関数をそのまま組み込みました）
def process_text(text):
    text = re.sub(r'[（\(].*?[）\)]', '', text)
    text = re.sub(r'[、,]', ' ', text)
    text = re.sub(r'([!?])(?![!?])', lambda m: {'!': '！', '?': '？'}[m.group(0)], text)
    text = re.sub(r'…{2,}|…{1,}', '…', text.replace('...', '…'))
    text = re.sub(r'[Ａ-Ｚａ-ｚ０-９]', lambda m: chr(ord(m.group(0)) - 0xFEE0), text)
    text = re.sub(r'\d+', lambda m: chr(ord(m.group(0)) + 0xFEE0) if len(m.group(0)) == 1 else m.group(0), text)
    text = re.sub(r'[ 　]{2,}', ' ', text)
    text = text.replace('【', '@@@【').replace('】', '】@@@')
    text = text.replace('。', '。@@@')
    text = text.replace('。', '　')
    return text

st.title("字幕作成ツール")
st.write("Wordファイルをアップロードすると、字幕用に最適化されたファイルを生成します。")

uploaded_file = st.file_uploader("input.docx を選択してください", type=["docx"])

if uploaded_file:
    # ファイル処理
    doc = Document(uploaded_file)
    new_doc = Document()
    
    for para in doc.paragraphs:
        if not para.text.strip(): continue
        processed_text = process_text(para.text)
        
        for segment in processed_text.split('@@@'):
            if not segment.strip(): continue
            if '【' in segment:
                new_doc.add_paragraph()
            
            all_chars = list(segment)
            current_line = []
            count = 0
            for char in all_chars:
                current_line.append(char)
                if char not in [' ', '　']:
                    count += 1
                if count >= 17:
                    p = new_doc.add_paragraph("".join(current_line))
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    current_line = []
                    count = 0
            if current_line:
                p = new_doc.add_paragraph("".join(current_line))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ダウンロード用バッファ
    buffer = io.BytesIO()
    new_doc.save(buffer)
    
    st.success("変換が完了しました！")
    st.download_button(
        label="変換後のファイルをダウンロード",
        data=buffer.getvalue(),
        file_name="output_subtitle.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )